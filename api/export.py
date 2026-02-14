import io
import re

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import PlainTextResponse, JSONResponse, StreamingResponse
from sqlalchemy.orm import Session, joinedload

from database import get_db
from models import Meeting, Segment
from models.action import Action, ActionResult

router = APIRouter(prefix="/api", tags=["export"])

UNKNOWN_SPEAKER = "Okand"


def _safe_filename(name: str) -> str:
    """Sanitize a string for use in Content-Disposition filename."""
    safe = re.sub(r'["\\/:<>|?*\x00-\x1f]', '_', name)
    return safe.strip('. ') or "export"


def format_srt_time(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def format_vtt_time(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


def format_timestamp_short(seconds: float) -> str:
    """Format seconds to MM:SS or H:MM:SS."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    if h > 0:
        return f"{h}:{m:02d}:{s:02d}"
    return f"{m}:{s:02d}"


@router.get("/meetings/{meeting_id}/export")
def export_meeting(
    meeting_id: str,
    format: str = "srt",
    db: Session = Depends(get_db),
):
    meeting = db.query(Meeting).filter(Meeting.id == meeting_id).first()
    if not meeting:
        raise HTTPException(404, "Meeting not found")

    segments = (
        db.query(Segment)
        .options(joinedload(Segment.speaker))
        .filter(Segment.meeting_id == meeting_id)
        .order_by(Segment.order)
        .all()
    )

    if format == "srt":
        return _export_srt(meeting, segments)
    elif format == "vtt":
        return _export_vtt(meeting, segments)
    elif format == "txt":
        return _export_txt(meeting, segments)
    elif format == "json":
        return _export_json(meeting, segments)
    elif format == "md":
        return _export_md(meeting, segments)
    elif format == "docx":
        return _export_docx(meeting, segments)
    elif format == "pdf":
        return _export_pdf(meeting, segments)
    else:
        raise HTTPException(400, f"Unknown format: {format}")


def _export_srt(meeting: Meeting, segments: list[Segment]) -> PlainTextResponse:
    lines = []
    for i, seg in enumerate(segments, 1):
        speaker = seg.speaker.display_name if seg.speaker else UNKNOWN_SPEAKER
        lines.append(str(i))
        lines.append(f"{format_srt_time(seg.start_time)} --> {format_srt_time(seg.end_time)}")
        lines.append(f"[{speaker}] {seg.text}")
        lines.append("")

    content = "\n".join(lines)
    return PlainTextResponse(
        content,
        headers={
            "Content-Disposition": f'attachment; filename="{_safe_filename(meeting.title)}.srt"',
        },
    )


def _export_vtt(meeting: Meeting, segments: list[Segment]) -> PlainTextResponse:
    lines = ["WEBVTT", ""]
    for seg in segments:
        speaker = seg.speaker.display_name if seg.speaker else UNKNOWN_SPEAKER
        lines.append(f"{format_vtt_time(seg.start_time)} --> {format_vtt_time(seg.end_time)}")
        lines.append(f"<v {speaker}>{seg.text}")
        lines.append("")

    content = "\n".join(lines)
    return PlainTextResponse(
        content,
        headers={
            "Content-Disposition": f'attachment; filename="{_safe_filename(meeting.title)}.vtt"',
        },
    )


def _export_txt(meeting: Meeting, segments: list[Segment]) -> PlainTextResponse:
    lines = []
    current_speaker = None
    for seg in segments:
        speaker = seg.speaker.display_name if seg.speaker else UNKNOWN_SPEAKER
        if speaker != current_speaker:
            if lines:
                lines.append("")
            lines.append(f"{speaker}:")
            current_speaker = speaker
        lines.append(f"  {seg.text}")

    content = "\n".join(lines)
    return PlainTextResponse(
        content,
        headers={
            "Content-Disposition": f'attachment; filename="{_safe_filename(meeting.title)}.txt"',
        },
    )


def _export_json(meeting: Meeting, segments: list[Segment]) -> JSONResponse:
    data = {
        "meeting": {
            "title": meeting.title,
            "duration": meeting.duration,
        },
        "segments": [
            {
                "start": seg.start_time,
                "end": seg.end_time,
                "speaker": seg.speaker.display_name if seg.speaker else UNKNOWN_SPEAKER,
                "text": seg.text,
            }
            for seg in segments
        ],
    }
    return JSONResponse(
        data,
        headers={
            "Content-Disposition": f'attachment; filename="{_safe_filename(meeting.title)}.json"',
        },
    )


def _export_md(meeting: Meeting, segments: list[Segment]) -> PlainTextResponse:
    lines = [f"# {meeting.title}", ""]
    if meeting.duration:
        lines.append(f"**Duration:** {format_timestamp_short(meeting.duration)}")
        lines.append("")

    current_speaker = None
    for seg in segments:
        speaker = seg.speaker.display_name if seg.speaker else UNKNOWN_SPEAKER
        if speaker != current_speaker:
            lines.append("")
            ts = format_timestamp_short(seg.start_time)
            lines.append(f"### {speaker} [{ts}]")
            lines.append("")
            current_speaker = speaker
        lines.append(seg.text)

    content = "\n".join(lines)
    return PlainTextResponse(
        content,
        headers={
            "Content-Disposition": f'attachment; filename="{_safe_filename(meeting.title)}.md"',
            "Content-Type": "text/markdown; charset=utf-8",
        },
    )


def _export_docx(meeting: Meeting, segments: list[Segment]) -> StreamingResponse:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(meeting.title, level=1)

    if meeting.duration:
        p = doc.add_paragraph()
        run = p.add_run(f"Duration: {format_timestamp_short(meeting.duration)}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

    current_speaker = None
    for seg in segments:
        speaker = seg.speaker.display_name if seg.speaker else UNKNOWN_SPEAKER
        if speaker != current_speaker:
            ts = format_timestamp_short(seg.start_time)
            doc.add_heading(f"{speaker} [{ts}]", level=3)
            current_speaker = speaker
        doc.add_paragraph(seg.text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{_safe_filename(meeting.title)}.docx"',
        },
    )


def _export_pdf(meeting: Meeting, segments: list[Segment]) -> StreamingResponse:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=20 * mm, bottomMargin=20 * mm)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "MeetingTitle",
        parent=styles["Title"],
        fontSize=18,
        spaceAfter=6,
    )
    speaker_style = ParagraphStyle(
        "SpeakerHeading",
        parent=styles["Heading3"],
        fontSize=11,
        textColor=HexColor("#6366f1"),
        spaceBefore=12,
        spaceAfter=4,
    )
    text_style = ParagraphStyle(
        "SegmentText",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=2,
    )
    meta_style = ParagraphStyle(
        "Meta",
        parent=styles["Normal"],
        fontSize=9,
        textColor=HexColor("#888888"),
        spaceAfter=8,
    )

    story = []
    story.append(Paragraph(meeting.title, title_style))

    if meeting.duration:
        story.append(Paragraph(f"Duration: {format_timestamp_short(meeting.duration)}", meta_style))

    current_speaker = None
    for seg in segments:
        speaker = seg.speaker.display_name if seg.speaker else UNKNOWN_SPEAKER
        if speaker != current_speaker:
            ts = format_timestamp_short(seg.start_time)
            story.append(Paragraph(f"{speaker} [{ts}]", speaker_style))
            current_speaker = speaker
        # Escape XML special chars for reportlab
        safe_text = (seg.text
                     .replace("&", "&amp;")
                     .replace("<", "&lt;")
                     .replace(">", "&gt;"))
        story.append(Paragraph(safe_text, text_style))

    doc.build(story)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{_safe_filename(meeting.title)}.pdf"',
        },
    )


# --- Action result export ---

@router.get("/actions/results/{result_id}/export")
def export_action_result(
    result_id: str,
    format: str = "txt",
    db: Session = Depends(get_db),
):
    result = db.query(ActionResult).filter(ActionResult.id == result_id).first()
    if not result:
        raise HTTPException(404, "Action result not found")
    if not result.result_text:
        raise HTTPException(400, "Action result has no content")

    action = db.query(Action).filter(Action.id == result.action_id).first()
    action_name = action.name if action else "Action"

    meeting = db.query(Meeting).filter(Meeting.id == result.meeting_id).first()
    meeting_title = meeting.title if meeting else "Meeting"

    filename = _safe_filename(f"{meeting_title} - {action_name}")

    if format == "txt":
        return PlainTextResponse(
            result.result_text,
            headers={"Content-Disposition": f'attachment; filename="{filename}.txt"'},
        )
    elif format == "md":
        md_content = f"# {action_name}\n\n**Meeting:** {meeting_title}\n\n---\n\n{result.result_text}"
        return PlainTextResponse(
            md_content,
            headers={
                "Content-Disposition": f'attachment; filename="{filename}.md"',
                "Content-Type": "text/markdown; charset=utf-8",
            },
        )
    elif format == "docx":
        return _export_action_docx(result.result_text, action_name, meeting_title, filename)
    elif format == "pdf":
        return _export_action_pdf(result.result_text, action_name, meeting_title, filename)
    else:
        raise HTTPException(400, f"Unknown format: {format}")


def _export_action_docx(text: str, action_name: str, meeting_title: str, filename: str) -> StreamingResponse:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(action_name, level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"Meeting: {meeting_title}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    for line in text.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.strip():
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'},
    )


def _export_action_pdf(text: str, action_name: str, meeting_title: str, filename: str) -> StreamingResponse:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=20 * mm, bottomMargin=20 * mm)
    styles = getSampleStyleSheet()

    title_style = styles["Title"]
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=9, textColor=HexColor("#888888"), spaceAfter=8)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=4)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13, spaceBefore=10, spaceAfter=4)
    h3_style = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=11, spaceBefore=8, spaceAfter=4)

    story = []
    story.append(Paragraph(action_name, title_style))
    story.append(Paragraph(f"Meeting: {meeting_title}", meta_style))

    for line in text.split("\n"):
        safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if line.startswith("### "):
            story.append(Paragraph(safe[4:], h3_style))
        elif line.startswith("## "):
            story.append(Paragraph(safe[3:], h2_style))
        elif line.startswith("# "):
            story.append(Paragraph(safe[2:], title_style))
        elif safe.strip():
            story.append(Paragraph(safe, body_style))
        else:
            story.append(Spacer(1, 4))

    doc.build(story)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'},
    )
