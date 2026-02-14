import io
import logging
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse
from sqlalchemy.orm import Session, joinedload

from database import get_db
from models import Meeting, Segment, Speaker
from services.llm_service import LLMService
from model_config import get_model_config

log = logging.getLogger(__name__)
router = APIRouter(prefix="/api", tags=["protocol"])

MAX_TRANSCRIPT_CHARS = 15000

PROTOCOL_PROMPT = """Du ar en professionell protokollforsare for svensk offentlig sektor.
Skapa ett formellt motesprotokoll baserat pa transkriberingen nedan.

Protokollet ska folja svensk standard for offentlig sektor med foljande struktur:

1. **Rubrik** - Protokoll fran [typ av mote]
2. **Datum och tid** - Om det framgar, annars anvand det du vet
3. **Narvarande** - Lista alla deltagare (anvand namnen fran transkriberingen)
4. **Dagordning** - Identifiera dagordningspunkter fran vad som diskuterades
5. **Paragrafpunkter** - Numrerade paragrafer (§1, §2, etc.) med:
   - Rubrik for varje punkt
   - Kort sammanfattning av diskussionen
   - Eventuella BESLUT markerade tydligt
   - Eventuella ATGARDSPUNKTER med ansvarig person
6. **Justering** - Plats for underskrifter

Skriv pa korrekt svenska. Var koncis men fullstandig.
Markera beslut med "BESLUT:" och atgardspunkter med "ATGARD:".
Svara med ren text (inte JSON), formaterad med markdown."""


@router.post("/meetings/{meeting_id}/generate-protocol")
def generate_protocol(meeting_id: str, db: Session = Depends(get_db)):
    """Generate a formal Swedish meeting protocol (protokoll) using LLM."""
    meeting = db.query(Meeting).filter(Meeting.id == meeting_id).first()
    if not meeting:
        raise HTTPException(404, "Meeting not found")
    if meeting.status.value != "completed":
        raise HTTPException(400, "Meeting must be completed first")

    segments = (
        db.query(Segment)
        .options(joinedload(Segment.speaker))
        .filter(Segment.meeting_id == meeting_id)
        .order_by(Segment.order)
        .all()
    )
    if not segments:
        raise HTTPException(400, "No transcript segments found")

    speakers = db.query(Speaker).filter(Speaker.meeting_id == meeting_id).all()
    speaker_map = {s.id: s.display_name or s.label for s in speakers}
    speaker_names = [s.display_name or s.label for s in speakers if s.label != "UNKNOWN"]

    lines = []
    for seg in segments:
        speaker = speaker_map.get(seg.speaker_id, "Okand") if seg.speaker_id else "Okand"
        ts = f"{int(seg.start_time // 60)}:{int(seg.start_time % 60):02d}"
        lines.append(f"[{ts}] [{speaker}]: {seg.text}")

    transcript_text = "\n".join(lines)
    if len(transcript_text) > MAX_TRANSCRIPT_CHARS:
        transcript_text = transcript_text[:MAX_TRANSCRIPT_CHARS] + "\n\n[...transkribering trunkerad...]"

    preset = get_model_config().get_model_for_task("actions")
    llm = LLMService(preset=preset)

    date_str = meeting.created_at.strftime("%Y-%m-%d") if meeting.created_at else ""
    duration_str = ""
    if meeting.duration:
        mins = int(meeting.duration // 60)
        duration_str = f"{mins} minuter"

    messages = [
        {"role": "system", "content": PROTOCOL_PROMPT},
        {"role": "user", "content": (
            f"Motestitel: {meeting.title}\n"
            f"Datum: {date_str}\n"
            f"Langd: {duration_str}\n"
            f"Deltagare: {', '.join(speaker_names)}\n\n"
            f"Transkribering:\n{transcript_text}"
        )},
    ]

    protocol_text = llm._call(messages, max_tokens=4000)
    return {"protocol_text": protocol_text}


def _add_inline_runs(paragraph, text, base_color=None):
    """Parse inline markdown (bold, italic, bold+italic) and add as runs."""
    import re
    from docx.shared import RGBColor

    # Split on bold+italic, bold, or italic markers
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        if base_color:
            run.font.color.rgb = base_color


def _md_to_docx(doc, protocol_text):
    """Convert markdown protocol text to formatted DOCX paragraphs."""
    import re
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = protocol_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = doc.styles["Normal"].font.size
            run = p.add_run("─" * 50)
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(180, 180, 180)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip("*").strip()
            p = doc.add_heading(heading_text, level=level)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # § headings (treat as h2)
        if re.match(r'^§\s*\d+', stripped):
            doc.add_heading(stripped, level=2)
            i += 1
            continue

        # Bullet list item
        m_bullet = re.match(r'^[-*]\s+(.*)', stripped)
        if m_bullet:
            content = m_bullet.group(1)
            p = doc.add_paragraph(style="List Bullet")
            if "BESLUT:" in content:
                _add_inline_runs(p, content, base_color=RGBColor(0, 100, 0))
            elif "ATGARD:" in content or "ÅTGÄRD:" in content:
                _add_inline_runs(p, content, base_color=RGBColor(0, 0, 150))
            else:
                _add_inline_runs(p, content)
            i += 1
            continue

        # Numbered list item
        m_num = re.match(r'^\d+\.\s+(.*)', stripped)
        if m_num:
            content = m_num.group(1)
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, content)
            i += 1
            continue

        # BESLUT / ÅTGÄRD highlight lines
        if "BESLUT:" in stripped:
            p = doc.add_paragraph()
            _add_inline_runs(p, stripped, base_color=RGBColor(0, 100, 0))
            for run in p.runs:
                run.bold = True
            i += 1
            continue

        if "ATGARD:" in stripped or "ÅTGÄRD:" in stripped:
            p = doc.add_paragraph()
            _add_inline_runs(p, stripped, base_color=RGBColor(0, 0, 150))
            for run in p.runs:
                run.bold = True
            i += 1
            continue

        # Regular paragraph — collect continuation lines
        para_lines = [stripped]
        while i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            # Stop at empty lines, headings, lists, rules
            if not next_line or next_line.startswith("#") or next_line.startswith("- ") or next_line.startswith("* ") or re.match(r'^\d+\.\s', next_line) or re.match(r'^§\s*\d+', next_line) or re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', next_line):
                break
            para_lines.append(next_line)
            i += 1

        full_text = " ".join(para_lines)
        p = doc.add_paragraph()
        _add_inline_runs(p, full_text)
        i += 1


@router.post("/meetings/{meeting_id}/export-protocol")
def export_protocol_docx(meeting_id: str, body: dict, db: Session = Depends(get_db)):
    """Export a protocol as DOCX with formal formatting."""
    from docx import Document
    from docx.shared import Pt
    import re

    protocol_text = body.get("protocol_text", "")
    if not protocol_text:
        raise HTTPException(400, "No protocol text provided")

    meeting = db.query(Meeting).filter(Meeting.id == meeting_id).first()
    if not meeting:
        raise HTTPException(404, "Meeting not found")

    doc = Document()

    # Style defaults
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    _md_to_docx(doc, protocol_text)

    # Add signature lines
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph("_" * 40)
    p.add_run("\nOrdforande")
    doc.add_paragraph()
    p = doc.add_paragraph("_" * 40)
    p.add_run("\nJusterare")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = re.sub(r'["\\/:<>|?*\x00-\x1f]', '_', meeting.title)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="Protokoll - {safe_title}.docx"',
        },
    )
